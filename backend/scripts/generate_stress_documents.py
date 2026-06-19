"""Generate deterministic large documents for local stress simulations."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

import requests
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


ROOT_DIR = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT_DIR / "backend" / "runtime" / "stress_documents"
IRPF_2026_URL = (
    "https://www.gov.br/receitafederal/pt-br/centrais-de-conteudo/"
    "publicacoes/perguntas-e-respostas/dirpf/"
    "p-r-irpf-2026-v1-00-2026-04-23.pdf"
)

PROFESSIONAL_PARAGRAPH = (
    "CASO ATLAS 2026. Contrato de prestacao de servicos entre Alfa Ltda. e "
    "Beta S.A., valor mensal de R$ 18.750,00, prazo de notificacao de 15 dias "
    "e foro de Curitiba. A empresa possui empregados registrados no eSocial, "
    "apuracao previdenciaria na DCTFWeb e escrituracao contabil regular. "
    "Revisar provas, obrigacoes acessorias, prazos e conciliacao contabil. "
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def create_large_text(path: Path, target_mb: int = 12) -> None:
    target_size = target_mb * 1024 * 1024
    with path.open("wb") as handle:
        handle.write(b"MARCADOR-INICIO-ATLAS\n")
        encoded = (PROFESSIONAL_PARAGRAPH + "\n").encode("utf-8")
        while handle.tell() + len(encoded) < target_size - 64:
            handle.write(encoded)
        handle.write(b"\nMARCADOR-FINAL-ATLAS\n")


def create_oversized_text(path: Path, target_mb: int = 51) -> None:
    with path.open("wb") as handle:
        handle.write(b"ARQUIVO ACIMA DO LIMITE\n")
        handle.truncate(target_mb * 1024 * 1024)


def create_pdf(path: Path, pages: int) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4, pageCompression=1)
    width, height = A4
    for page_number in range(1, pages + 1):
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(48, height - 52, f"Relatorio profissional Atlas - pagina {page_number}")
        pdf.setFont("Helvetica", 9)
        y = height - 78
        for line_number in range(1, 38):
            text = (
                f"{line_number:02d}. {PROFESSIONAL_PARAGRAPH[:112]} "
                f"Ref. PAG-{page_number:04d}-LIN-{line_number:02d}"
            )
            pdf.drawString(48, y, text[:120])
            y -= 18
        pdf.showPage()
    pdf.save()


def create_docx(path: Path, paragraphs: int = 9000) -> None:
    document = Document()
    document.add_heading("Dossie profissional Atlas", level=1)
    for index in range(1, paragraphs + 1):
        document.add_paragraph(
            f"{index:05d} - {PROFESSIONAL_PARAGRAPH} REF-DOCX-{index:05d}"
        )
    document.save(path)


def create_rtf(path: Path, target_mb: int = 5) -> None:
    target_size = target_mb * 1024 * 1024
    with path.open("wb") as handle:
        handle.write(b"{\\rtf1\\ansi\\deff0 ")
        line = (
            PROFESSIONAL_PARAGRAPH.replace("\\", "\\\\")
            + r" Referencia RTF-ATLAS\par "
        ).encode("cp1252", errors="replace")
        while handle.tell() + len(line) < target_size - 2:
            handle.write(line)
        handle.write(b"}")


def download_official_irpf(path: Path) -> str:
    try:
        response = requests.get(IRPF_2026_URL, timeout=90)
        response.raise_for_status()
        path.write_bytes(response.content)
        return "downloaded"
    except Exception as exc:
        return f"unavailable: {exc}"


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    artifacts = {
        "large_txt": OUTPUT_DIR / "atlas_12mb.txt",
        "oversized_txt": OUTPUT_DIR / "acima_limite_51mb.txt",
        "large_pdf": OUTPUT_DIR / "atlas_250_paginas.pdf",
        "page_limit_pdf": OUTPUT_DIR / "acima_limite_501_paginas.pdf",
        "large_docx": OUTPUT_DIR / "atlas_9000_paragrafos.docx",
        "large_rtf": OUTPUT_DIR / "atlas_5mb.rtf",
        "disguised_pdf": OUTPUT_DIR / "executavel_disfarcado.pdf",
        "official_irpf_2026": OUTPUT_DIR / "receita_irpf_2026.pdf",
    }

    create_large_text(artifacts["large_txt"])
    create_oversized_text(artifacts["oversized_txt"])
    create_pdf(artifacts["large_pdf"], pages=250)
    create_pdf(artifacts["page_limit_pdf"], pages=501)
    create_docx(artifacts["large_docx"])
    create_rtf(artifacts["large_rtf"])
    artifacts["disguised_pdf"].write_bytes(b"MZ\x90\x00FAKE-EXECUTABLE")
    official_status = download_official_irpf(artifacts["official_irpf_2026"])

    manifest = {
        "official_irpf_status": official_status,
        "artifacts": {
            name: {
                "path": str(path),
                "size_bytes": path.stat().st_size if path.exists() else 0,
                "sha256": sha256(path) if path.exists() else None,
            }
            for name, path in artifacts.items()
        },
    }
    manifest_path = OUTPUT_DIR / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
